from datetime import datetime
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
import lancedb
from sentence_transformers import SentenceTransformer
from rag_core.semantic_chunking import embedding_semantic_chunk_text
from .models import Document
from .word_chunking import word_chunk_text
from .embeddings import compute_embeddings
from .config import DB_DIR, TABLE_NAME, CHUNK_SIZE_WORDS, CHUNK_OVERLAP_WORDS, CHUNK_STRATEGY
from pypdf import PdfReader 
from docx import Document as DocxDocument     # <-- NEW

# ---- NEW: supported extensions & helpers ----

SUPPORTED_EXTENSIONS = (".txt", ".md", ".pdf", ".docx")

def get_db() -> lancedb.DBConnection:
    return lancedb.connect(DB_DIR)


def open_or_create_table(db: lancedb.DBConnection) -> lancedb.table.LanceTable:
    """
    Open the LanceDB table if it exists, otherwise create an empty one.
    We won't drop the table here; this is for incremental indexing.
    """
    if TABLE_NAME in db.table_names():
        return db.open_table(TABLE_NAME)

    # Create an empty table schema by creating with no rows first time.
    # We'll create it when we get first records.
    return None


def get_next_doc_version(table, doc_key: str) -> int:
    """
    Look at existing rows for this doc_key in LanceDB and decide the next doc_version.
    If doc_key does not exist yet, return 1.
    """
    # We only need doc_key + doc_version to compute max
    df = table.to_pandas()[["doc_key", "doc_version"]]
    if df.empty:
        return 1

    rows = df[df["doc_key"] == doc_key]
    if rows.empty:
        return 1

    return int(rows["doc_version"].max()) + 1


def index_single_document_incremental(
    path: Path,
    model: SentenceTransformer,
    db: lancedb.DBConnection,
    table: lancedb.table.LanceTable | None,
) -> Dict[str, Any]:
    """
    Incrementally index a single document into LanceDB.

    Uses your existing build_document_from_path + chunking + compute_embeddings.
    Does NOT drop the table; it only adds records.

    Returns a small summary:
      - doc_key
      - doc_version
      - doc_type
      - num_chunks
      - source
    """
    if not path.exists() or not path.is_file():
        print(f"[INDEX] Skipping missing/non-file path: {path}")
        return {
            "doc_key": path.stem,
            "doc_version": 0,
            "doc_type": "unknown",
            "num_chunks": 0,
            "source": str(path),
        }

    # If table is None, we haven't created it yet.
    doc_key = path.stem

    # If you still want build_document_from_path to control doc_type, created_at etc,
    # we call it AFTER we know the version.
    # First, compute next version for this doc_key from existing table data.
    next_version = get_next_doc_version(table, doc_key)

    doc = build_document_from_path(path, version=next_version)

    print(f"\n[INDEX] Incremental indexing: {path.name} as "
          f"doc_key={doc.doc_key}, version={doc.doc_version}")

    page_fragments = get_page_fragments(doc.path)
    print(f"  -> found {len(page_fragments)} page fragments")

    all_records: List[Dict[str, Any]] = []
    chunk_global_id = 0

    for page_number, page_text in page_fragments:
        if not page_text or not page_text.strip():
            continue

        if CHUNK_STRATEGY == "word":
            chunks = word_chunk_text(
                page_text,
                chunk_size=CHUNK_SIZE_WORDS,
                overlap=CHUNK_OVERLAP_WORDS,
            )
        elif CHUNK_STRATEGY == "semantic":
            chunks = embedding_semantic_chunk_text(
                page_text,
                model=model,
                target_chunk_size_words=CHUNK_SIZE_WORDS,
                min_chunk_size_words=int(CHUNK_SIZE_WORDS * 0.4),
                similarity_percentile=30.0,
            )
        else:
            raise ValueError(f"Unknown chunking strategy: {CHUNK_STRATEGY}")

        print(f"    page={page_number} -> produced {len(chunks)} chunks")

        if not chunks:
            continue

        embeddings = compute_embeddings(model, chunks)

        for i, (chunk_text, emb) in enumerate(zip(chunks, embeddings)):
            all_records.append(
                {
                    "doc_key": doc.doc_key,
                    "doc_version": doc.doc_version,
                    "page_number": page_number,
                    "chunk_id": chunk_global_id,
                    "doc_type": doc.doc_type,
                    "source": str(doc.path),
                    "created_at": doc.created_at.isoformat(),
                    "text": chunk_text,
                    "embedding": emb.tolist(),
                }
            )
            chunk_global_id += 1

    if not all_records:
        print(f"[INDEX] No chunks produced for {path}, skipping insert.")
        return {
            "doc_key": doc.doc_key,
            "doc_version": doc.doc_version,
            "doc_type": doc.doc_type,
            "num_chunks": 0,
            "source": str(doc.path),
        }

    # Create table if needed, otherwise append
    if table is None:
        print(f"[INDEX] Table '{TABLE_NAME}' does not exist. Creating with first document...")
        table = db.create_table(TABLE_NAME, data=all_records)
    else:
        print(f"[INDEX] Appending {len(all_records)} chunks to existing table '{TABLE_NAME}'...")
        table.add(all_records)

    print(f"[INDEX] Completed {path.name}: {len(all_records)} chunks added.")

    summary = {
        "doc_key": doc.doc_key,
        "doc_version": doc.doc_version,
        "doc_type": doc.doc_type,
        "num_chunks": len(all_records),
        "source": str(doc.path),
    }
    return summary


def iter_source_files(data_dir: Path):
    """
    Yield all files in data_dir with supported extensions (.txt, .md, ...).
    """
    for ext in SUPPORTED_EXTENSIONS:
        for path in data_dir.glob(f"*{ext}"):
            yield path


def load_pdf_pages(path: Path) -> List[Tuple[int, str]]:
    """
    Return list of (page_number, text) for the PDF.
    page_number is 1-based.
    """
    reader = PdfReader(str(path))
    pages: List[Tuple[int, str]] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if not text:
            continue
        pages.append((i + 1, text))
    return pages


def load_docx_text(path: Path) -> str:
    """
    Load text from a DOCX file by concatenating all paragraph texts.
    No page numbers are available at this level.
    """
    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def get_page_fragments(path: Path) -> List[Tuple[Optional[int], str]]:
    """
    Normalize content from different file types into a list of (page_number, text).

    - For PDFs: list of real pages.
    - For txt/md: single 'page' with page_number=None.
    """
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf_pages(path)

    if suffix in (".txt", ".md"):
        text = path.read_text(encoding="utf-8")
        return [(None, text)]

    if suffix == ".docx":
        text = load_docx_text(path)
        return [(None, text)]

    raise ValueError(f"Unsupported file extension: {suffix} for {path}")
    

def _infer_doc_type(path: Path) -> str:
    lower_name = path.name.lower()
    if "policy" in lower_name:
        return "policy"
    if "tutorial" in lower_name:
        return "tutorial"
    return "general"


def build_document_from_path(path: Path, version: int) -> Document:
    doc_key = path.stem
    doc_type = _infer_doc_type(path)
    return Document(
        doc_key=doc_key,
        doc_version=version,
        path=path,
        doc_type=doc_type,
        created_at=datetime.utcnow(),
    )


def rebuild_all_documents_versioned(
    data_dir: Path,
    model: SentenceTransformer,
    db: lancedb.DBConnection,
    table_name: str,
    version: int,
):
    """
    Index all supported docs (.txt, .md) as a specific version into the same table.
    If table doesn't exist, create it. Otherwise append.
    """
    files = sorted(iter_source_files(data_dir))
    if not files:
        raise RuntimeError(f"No supported files ({SUPPORTED_EXTENSIONS}) found in {data_dir}")

    all_records: List[Dict] = []

    for path in files:
        print(f"\nIndexing document: {path.name} as version {version}")
        doc = build_document_from_path(path, version=version)

        # NEW: iterate page fragments (per-page for PDF, single for txt/md)
        page_fragments = get_page_fragments(doc.path)
        print(f"  -> found {len(page_fragments)} page fragments")

        for page_number, page_text in page_fragments:
            if CHUNK_STRATEGY == "word":
                chunks = word_chunk_text(
                    page_text,
                    chunk_size=CHUNK_SIZE_WORDS,
                    overlap=CHUNK_OVERLAP_WORDS,
                )
            elif CHUNK_STRATEGY == "semantic":
                chunks = embedding_semantic_chunk_text(
                    page_text,
                    model=model,
                    target_chunk_size_words=CHUNK_SIZE_WORDS,
                    min_chunk_size_words=int(CHUNK_SIZE_WORDS * 0.4),
                    similarity_percentile=30.0
                )
            else:
                raise ValueError(f"Unknown chunking strategy: {CHUNK_STRATEGY}")
            print(f"    page={page_number} -> produced {len(chunks)} chunks")

            if not chunks:
                continue

            embeddings = compute_embeddings(model, chunks)

            for i, (chunk_text, emb) in enumerate(zip(chunks, embeddings)):
                all_records.append(
                    {
                        "doc_key": doc.doc_key,
                        "doc_version": doc.doc_version,
                        "page_number": page_number,          # <-- NEW
                        "chunk_id": i,
                        "doc_type": doc.doc_type,
                        "source": str(doc.path),
                        "created_at": doc.created_at.isoformat(),
                        "text": chunk_text,
                        "embedding": emb.tolist(),
                    }
                )

    if table_name in db.table_names():
        print(f"Table '{table_name}' exists. Appending records...")
        db.drop_table(table_name)

    print(f"Creating table '{table_name}' with initial records...")
    table = db.create_table(table_name, data=all_records)

    return table

def index_paths_incremental(
    paths: List[Path],
    model: SentenceTransformer,
) -> List[Dict[str, Any]]:
    """
    Incrementally index the given paths into LanceDB.

    - Reuses your existing chunking + embeddings logic.
    - Does NOT drop or recreate the table.
    - Assigns doc_version per doc_key based on existing rows.

    Returns a list of per-document summaries.
    """
    db = get_db()
    table = open_or_create_table(db)

    summaries: List[Dict[str, Any]] = []

    for path in paths:
        summary = index_single_document_incremental(
            path=path,
            model=model,
            db=db,
            table=table,
        )
        summaries.append(summary)

    return summaries